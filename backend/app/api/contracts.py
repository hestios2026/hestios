"""Contract generation — Arbeitsvertrag DOCX download."""
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from datetime import date
from dateutil.relativedelta import relativedelta
from urllib.parse import quote
import unicodedata
import io

from app.core.database import get_db
from app.models.employee import Employee
from app.models.user import User
from app.api.auth import get_current_user

router = APIRouter(prefix="/contracts", tags=["contracts"])

MONTHS_DE = ["Januar","Februar","März","April","Mai","Juni","Juli","August","September","Oktober","November","Dezember"]

CONTRACT_TYPE_DE = {
    "unbefristet": "Unbefristet",
    "befristet":   "Befristet",
    "minijob":     "Minijob",
}


def _ascii_filename(s: str) -> str:
    """Remove diacritics and non-latin-1 characters for Content-Disposition filename= fallback."""
    normalized = unicodedata.normalize("NFD", s)
    return "".join(c for c in normalized if unicodedata.category(c) != "Mn" and ord(c) < 256)


def _fmt_date(d) -> str:
    if not d:
        return "—"
    if isinstance(d, str):
        try:
            from datetime import datetime
            d = datetime.strptime(d, "%Y-%m-%d").date()
        except Exception:
            return str(d)
    return f"{d.day}. {MONTHS_DE[d.month-1]} {d.year}"


def _generate_arbeitsvertrag(emp: Employee) -> bytes:
    """Generate Arbeitsvertrag as DOCX bytes."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    def add_heading(text, level=1, center=False):
        p = doc.add_heading(text, level=level)
        p.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def add_field(label, value, doc=doc):
        p = doc.add_paragraph()
        run_lbl = p.add_run(f"{label}: ")
        run_lbl.bold = True
        run_lbl.font.size = Pt(11)
        run_val = p.add_run(str(value or "—"))
        run_val.font.size = Pt(11)
        return p

    def add_para(text, bold=False, size=11):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        return p

    def add_line():
        doc.add_paragraph("_" * 80)

    # ── Title ──────────────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ARBEITSVERTRAG")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)

    doc.add_paragraph()

    # ── Parteien ───────────────────────────────────────────────────────────────
    add_heading("§ 1 Vertragsparteien", level=2)
    add_para("Arbeitgeber:")
    add_field("Firma", "Hesti Rossmann GmbH")
    add_field("Adresse", "Austr. 6, 73230 Kirchheim unter Teck")
    doc.add_paragraph()
    add_para("Arbeitnehmer:")
    add_field("Name, Vorname",       f"{emp.nachname}, {emp.vorname}")
    add_field("Geburtsdatum",        _fmt_date(emp.geburtsdatum))
    add_field("Geburtsort",          emp.geburtsort or "—")
    add_field("Wohnhaft",            emp.adresse or "—")
    add_field("Staatsangehörigkeit", emp.nationalitaet or "—")
    add_field("Familienstand",       emp.familienstand or "—")
    if emp.kinder:
        add_field("Kinder", f"Ja ({emp.kinder_anzahl})")
    add_field("Konfession",          emp.konfession or "—")
    add_field("Schulabschluss",      emp.schulabschluss or "—")
    add_field("Berufsausbildung",    emp.berufsausbildung or "—")
    add_field("Beschäftigungsart",   emp.beschaeftigungsart or "Hauptbeschäftigung")

    doc.add_paragraph()

    # ── Beginn ─────────────────────────────────────────────────────────────────
    add_heading("§ 2 Beginn des Arbeitsverhältnisses", level=2)
    add_field("Arbeitsbeginn", _fmt_date(emp.arbeitsbeginn))

    ct_label = {"unbefristet": "Unbefristet", "befristet": "Befristet", "minijob": "Geringfügige Beschäftigung (Minijob)"}.get(emp.contract_type, emp.contract_type)
    add_field("Vertragsart", ct_label)

    if emp.contract_type == "befristet" and emp.befristung_bis:
        add_field("Befristet bis", _fmt_date(emp.befristung_bis))

    if emp.probezeit_end:
        add_field("Probezeit bis", _fmt_date(emp.probezeit_end))
    else:
        # Default: 6 months
        probezeit = None
        if emp.arbeitsbeginn:
            try:
                from datetime import datetime
                ab = emp.arbeitsbeginn if isinstance(emp.arbeitsbeginn, date) else datetime.strptime(str(emp.arbeitsbeginn), "%Y-%m-%d").date()
                probezeit = ab + relativedelta(months=6)
            except Exception:
                pass
        if probezeit:
            add_field("Probezeit bis", _fmt_date(probezeit))

    add_field("Wöchentliche Arbeitszeit", f"{emp.stunden_pro_woche or 40.0} Stunden")

    doc.add_paragraph()

    # ── Tätigkeit ──────────────────────────────────────────────────────────────
    add_heading("§ 3 Tätigkeit", level=2)
    add_field("Tätigkeit",       emp.taetigkeit or "Bauarbeiter")
    add_field("Erlernter Beruf", emp.erlernter_beruf or "—")
    add_para("Der Arbeitnehmer ist verpflichtet, auch andere zumutbare Tätigkeiten zu verrichten, "
             "sofern dies betrieblich erforderlich ist und dem Arbeitnehmer zumutbar ist.")

    doc.add_paragraph()

    # ── Vergütung ──────────────────────────────────────────────────────────────
    add_heading("§ 4 Vergütung", level=2)
    add_field("Lohngruppe",              str(emp.lohngruppe))
    add_field("Tariflohn",               f"€ {emp.tariflohn:.2f} / Stunde")
    add_field("Bauzuschlag",             f"€ {emp.bauzuschlag:.2f} / Stunde")
    add_field("Gesamttarifstundenlohn",  f"€ {(emp.tariflohn + emp.bauzuschlag):.2f} / Stunde")
    add_para("Die Vergütung richtet sich nach dem Bundesrahmentarifvertrag für das Baugewerbe (BRTV) "
             "in der jeweils gültigen Fassung. Die Auszahlung erfolgt monatlich am letzten Werktag des Monats "
             "auf das angegebene Bankkonto.")
    if emp.iban:
        add_field("IBAN",    emp.iban)
        add_field("BIC",     emp.bic or "—")
        add_field("Bank",    emp.kreditinstitut or "—")

    doc.add_paragraph()

    # ── Sozialversicherung ─────────────────────────────────────────────────────
    add_heading("§ 5 Sozialversicherung und Steuern", level=2)
    add_field("Krankenversicherung",     emp.krankenkasse or "—")
    add_field("Sozialversicherungs-Nr.", emp.sozialversicherungsnr or "—")
    add_field("Steuer-ID",               emp.steuer_id or "—")
    add_field("Steuerklasse",            str(emp.steuerklasse))
    add_field("Rentenversicherungs-Nr.", emp.rentenversicherungsnr or "—")
    if emp.personalnummer:
        add_field("Personalnummer",      emp.personalnummer)

    doc.add_paragraph()

    # ── Urlaub ─────────────────────────────────────────────────────────────────
    add_heading("§ 6 Urlaub", level=2)
    add_field("Jahresurlaubsanspruch", f"{emp.urlaubsanspruch_tage or 30} Arbeitstage")
    add_para("Der Urlaub ist rechtzeitig zu beantragen und richtet sich nach den betrieblichen Belangen "
             "sowie den Vorschriften des Bundesurlaubsgesetzes (BUrlG).")

    doc.add_paragraph()

    # ── Kündigung ──────────────────────────────────────────────────────────────
    add_heading("§ 7 Kündigung", level=2)
    add_para("Während der Probezeit kann das Arbeitsverhältnis beiderseits mit einer Frist von 2 Wochen "
             "gekündigt werden. Nach Ablauf der Probezeit gelten die gesetzlichen Kündigungsfristen "
             "gemäß § 622 BGB. Die Kündigung bedarf der Schriftform.")

    doc.add_paragraph()

    # ── SOKA-BAU ───────────────────────────────────────────────────────────────
    add_heading("§ 8 SOKA-BAU", level=2)
    add_para("Der Arbeitgeber ist Mitglied der Sozialkassen des Baugewerbes (SOKA-BAU). "
             "Urlaubsansprüche werden über die SOKA-BAU abgewickelt.")
    if emp.soka_bau_nr:
        add_field("SOKA-BAU-Nr.", emp.soka_bau_nr)

    doc.add_paragraph()

    # ── A1 / Entsendung ────────────────────────────────────────────────────────
    if emp.a1_bescheinigung_nr or emp.a1_gueltig_bis:
        add_heading("§ 9 Entsendung / A1-Bescheinigung", level=2)
        add_para("Der Arbeitnehmer wird im Rahmen der Entsendung (§ 4 SGB IV) tätig. "
                 "Die A1-Bescheinigung ist beizufügen.")
        if emp.a1_bescheinigung_nr:
            add_field("A1-Bescheinigung-Nr.", emp.a1_bescheinigung_nr)
        if emp.a1_gueltig_von:
            add_field("Gültig von", _fmt_date(emp.a1_gueltig_von))
        if emp.a1_gueltig_bis:
            add_field("Gültig bis", _fmt_date(emp.a1_gueltig_bis))
        doc.add_paragraph()

    # ── Sonstiges ─────────────────────────────────────────────────────────────
    add_heading("§ 10 Sonstige Vereinbarungen", level=2)
    add_para("Im Übrigen gelten die gesetzlichen Bestimmungen, insbesondere das Arbeitszeitgesetz (ArbZG), "
             "das Mindestlohngesetz (MiLoG) sowie die einschlägigen Tarifverträge des Baugewerbes.")

    doc.add_paragraph()

    # ── Unterschriften ─────────────────────────────────────────────────────────
    add_heading("Unterschriften", level=2)
    add_para(f"Kirchheim unter Teck, den {_fmt_date(date.today())}")
    doc.add_paragraph()

    tbl = doc.add_table(rows=3, cols=2)
    tbl.cell(0, 0).text = "Arbeitgeber"
    tbl.cell(0, 1).text = "Arbeitnehmer"
    tbl.cell(1, 0).text = ""
    tbl.cell(1, 1).text = ""
    tbl.cell(2, 0).text = "Hesti Rossmann GmbH"
    tbl.cell(2, 1).text = f"{emp.vorname} {emp.nachname}"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Routes ────────────────────────────────────────────────────────────────────

@router.get("/{employee_id}/arbeitsvertrag/")
def download_arbeitsvertrag(
    employee_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if current_user.role not in ("director", "projekt_leiter"):
        raise HTTPException(403, "Insufficient rights")
    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        raise HTTPException(404, "Employee not found")

    try:
        docx_bytes = _generate_arbeitsvertrag(emp)
    except ImportError:
        raise HTTPException(500, "python-docx not installed — run: pip install python-docx")

    ct_de    = CONTRACT_TYPE_DE.get(emp.contract_type or "unbefristet", "Arbeitsvertrag")
    fullname = f"{emp.nachname}_{emp.vorname}"
    filename = f"Arbeitsvertrag_{ct_de}_{fullname}.docx"

    # HTTP headers must be latin-1; use RFC 5987 filename* for full Unicode support
    filename_ascii   = _ascii_filename(filename)
    filename_encoded = quote(filename, safe="_-.")  # percent-encode for filename*

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": (
                f"attachment; filename=\"{filename_ascii}\"; "
                f"filename*=UTF-8''{filename_encoded}"
            )
        },
    )
